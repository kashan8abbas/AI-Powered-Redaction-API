from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw
import easyocr
import io
import numpy as np
import os
from rapidfuzz import fuzz

reader = easyocr.Reader(['en', 'ar'], verbose=False)

def redact_text_on_image(image: Image.Image, phrase: str, threshold: int = 70) -> Image.Image:
    draw = ImageDraw.Draw(image)
    results = reader.readtext(np.array(image))

    for (bbox, text, _) in results:
        similarity = fuzz.ratio(phrase.strip(), text.strip())  # More strict, full comparison
        if similarity >= threshold:
            print(f"🟥 Redacting from image: '{text}' (matched with: '{phrase}', similarity: {similarity})")
            top_left = tuple(map(int, bbox[0]))
            bottom_right = tuple(map(int, bbox[2]))
            draw.rectangle([top_left, bottom_right], fill="black")
    return image


def redact_docx(input_path: str, phrase: str, placeholder: str = "████████", threshold: int = 75) -> str:
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    original_doc = Document(input_path)
    new_doc = Document()

    print("🔍 Redacting paragraphs...")
    for para in original_doc.paragraphs:
        words = para.text.split()
        new_words = []
        redacted = False

        for word in words:
            similarity = fuzz.ratio(word.strip(), phrase.strip())
            if similarity >= threshold:
                print(f"🟥 Redacting word: '{word}' (matched with: '{phrase}', similarity: {similarity})")
                new_words.append(placeholder)
                redacted = True
            else:
                new_words.append(word)

        new_text = ' '.join(new_words)
        new_doc.add_paragraph(new_text)

    print("🖼️ Redacting images...")
    for shape in original_doc.inline_shapes:
        try:
            img_blob = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = original_doc.part.related_parts[img_blob]
            image_bytes = image_part.blob

            img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
            redacted_img = redact_text_on_image(img, phrase)

            out_io = io.BytesIO()
            redacted_img.save(out_io, format="PNG")
            out_io.seek(0)

            new_doc.add_picture(out_io, width=Inches(5))
        except Exception as e:
            print(f"⚠️ Failed to process image: {e}")
            continue

    base, ext = os.path.splitext(input_path)
    output_path = f"{base}_redacted{ext}"
    new_doc.save(output_path)
    print(f"✅ Saved redacted document to: {output_path}")
    return output_path

